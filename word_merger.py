#!/usr/bin/env python3
"""
Word Document Merger - Combine .docx files into one master document and export as .docx, .pdf, and/or .md.

Dependencies (install with pip):
    python-docx, tqdm, pypandoc, docx2pdf (optional, Windows/Mac), WeasyPrint (optional)
System dependencies (for headless conversion):
    Pandoc (for markdown / PDF fallback), LibreOffice (for PDF fallback)
"""

import argparse
import os
import shutil
import tempfile
from pathlib import Path
from typing import List, Optional, Literal

# ----------------------------------------------------------------------
# Third-party imports (graceful handling if missing)
# ----------------------------------------------------------------------
try:
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, Inches, RGBColor
except ImportError:
    sys.exit("python-docx is required. Install it with: pip install python-docx")

try:
    from tqdm import tqdm
except ImportError:

    import sys
    tqdm = None
    print("Tip: install tqdm for progress bars (pip install tqdm)")

try:
    import pypandoc
    PANDOC_AVAILABLE = True
except (ImportError, OSError):
    PANDOC_AVAILABLE = False
    print("Tip: install pypandoc and Pandoc (https://pandoc.org) for Markdown/PDF fallback.")

try:
    import docx2pdf
    DOCX2PDF_AVAILABLE = True
except ImportError:
    DOCX2PDF_AVAILABLE = False

try:
    import weasyprint
    WEASYPRINT_AVAILABLE = True
except ImportError:
    WEASYPRINT_AVAILABLE = False

# ----------------------------------------------------------------------
# Configuration
# ----------------------------------------------------------------------
OUTPUT_DIR_NAME = "Output"
OUTPUT_FILENAME = "Combined"

# ----------------------------------------------------------------------
# Helper: progress wrapper
# ----------------------------------------------------------------------
def optional_progress(iterable, desc="Working", **kwargs):
    """Wrap an iterable with tqdm if available."""
    if tqdm:
        return tqdm(iterable, desc=desc, **kwargs)
    return iterable

# ----------------------------------------------------------------------
# 1. Merge .docx files
# ----------------------------------------------------------------------
def add_page_break(doc):
    """Append a page break to the document."""
    doc.add_page_break()

def merge_documents(file_paths: List[Path], output_path: Path,
                    add_page_breaks: bool = True,
                    insert_toc: bool = True) -> None:
    """
    Merge multiple .docx files into one, preserving formatting.
    Optionally insert a Table of Contents as a Word field.
    """
    if not file_paths:
        raise ValueError("No .docx files to merge.")

    master = Document()

    # Insert TOC field at the very beginning if requested
    if insert_toc:
        paragraph = master.add_paragraph()
        run = paragraph.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar1)

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' TOC \\o "1-3" \\h \\z '
        run._r.append(instrText)

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar2)

        master.add_page_break()

    # Copy content from each source document
    for idx, path in enumerate(optional_progress(file_paths, desc="Merging documents")):
        sub_doc = Document(path)

        # If it's not the first document and page breaks are wanted, insert one
        if idx > 0 and add_page_breaks:
            master.add_page_break()

        # Copy all elements from the sub_document body into the master
        for element in sub_doc.element.body:
            if element.tag == qn('w:sectPr'):
                continue
            master.element.body.append(element)

        # --- FIX: safely copy header/footer of the first document only ---
        if idx == 0:
            try:
                # Copy header of the first section
                if sub_doc.sections[0].header:
                    for paragraph in sub_doc.sections[0].header.paragraphs:
                        master.sections[0].header.add_paragraph(
                            paragraph.text, paragraph.style
                        )
                # Copy footer of the first section
                if sub_doc.sections[0].footer:
                    for paragraph in sub_doc.sections[0].footer.paragraphs:
                        master.sections[0].footer.add_paragraph(
                            paragraph.text, paragraph.style
                        )
            except Exception as e:
                print(f"  Note: could not copy header/footer from {path.name} – {e}")
        # ---------------------------------------------------------------

    master.save(str(output_path))

# ----------------------------------------------------------------------
# 2. Convert merged .docx to PDF
# ----------------------------------------------------------------------
def convert_to_pdf(docx_path: Path, pdf_path: Path) -> bool:
    """
    Convert a .docx file to PDF using the best available method.
    Returns True on success.
    """
    # 1. Try docx2pdf (requires Microsoft Word on Windows/Mac)
    if DOCX2PDF_AVAILABLE and sys.platform in ("win32", "darwin"):
        try:
            docx2pdf.convert(str(docx_path), str(pdf_path))
            return True
        except Exception as e:
            print(f"  docx2pdf failed: {e}")

    # 2. Try LibreOffice in headless mode (cross-platform)
    if shutil.which("libreoffice") or shutil.which("soffice"):
        lo_bin = "libreoffice" if shutil.which("libreoffice") else "soffice"
        out_dir = str(pdf_path.parent)
        try:
            import subprocess
            subprocess.run(
                [lo_bin, "--headless", "--convert-to", "pdf", "--outdir", out_dir, str(docx_path)],
                check=True, capture_output=True, timeout=300
            )
            # LibreOffice names the output <original_name>.pdf; rename if needed
            lo_pdf = docx_path.with_suffix(".pdf")
            if lo_pdf != pdf_path and lo_pdf.exists():
                lo_pdf.rename(pdf_path)
            return True
        except Exception as e:
            print(f"  LibreOffice conversion failed: {e}")

    # 3. Fallback to Pandoc + WeasyPrint (or wkhtmltopdf)
    if PANDOC_AVAILABLE:
        pdf_engine = "weasyprint" if WEASYPRINT_AVAILABLE else None
        if not pdf_engine:
            # check for wkhtmltopdf
            if shutil.which("wkhtmltopdf"):
                pdf_engine = "wkhtmltopdf"
        if pdf_engine:
            try:
                pypandoc.convert_file(
                    str(docx_path), 'pdf',
                    outputfile=str(pdf_path),
                    extra_args=[f'--pdf-engine={pdf_engine}']
                )
                return True
            except Exception as e:
                print(f"  Pandoc PDF conversion failed: {e}")

    print("  No suitable PDF converter found. Install Word, LibreOffice, or Pandoc+WeasyPrint/wkhtmltopdf.")
    return False

# ----------------------------------------------------------------------
# 3. Convert merged .docx to Markdown
# ----------------------------------------------------------------------
def convert_to_markdown(docx_path: Path, md_path: Path, extract_media: bool = True) -> bool:
    """Convert .docx to Markdown using Pandoc, preserving images if requested."""
    if not PANDOC_AVAILABLE:
        print("  Pandoc is not available. Install pandoc and pypandoc to generate Markdown.")
        return False

    extra_args = []
    if extract_media:
        media_dir = md_path.parent / "media"
        media_dir.mkdir(exist_ok=True)
        extra_args.extend(["--extract-media", str(media_dir)])

    try:
        pypandoc.convert_file(
            str(docx_path), 'markdown',
            outputfile=str(md_path),
            extra_args=extra_args
        )
        return True
    except Exception as e:
        print(f"  Markdown conversion failed: {e}")
        return False

# ----------------------------------------------------------------------
# 4. Main workflow
# ----------------------------------------------------------------------
def run_workflow(input_folder: Path,
                 formats: List[Literal["docx", "pdf", "md"]],
                 sort_order: Literal["name", "date"] = "name",
                 add_toc: bool = True,
                 page_breaks: bool = True):
    """Execute the full merging and conversion pipeline."""
    input_folder = input_folder.resolve()
    if not input_folder.is_dir():
        raise NotADirectoryError(f"Folder not found: {input_folder}")

    # Gather .docx files
    docx_files = sorted(
        [f for f in input_folder.glob("*.docx") if f.is_file()],
        key=lambda p: p.name.lower() if sort_order == "name" else p.stat().st_mtime
    )

    if not docx_files:
        print(f"No .docx files found in {input_folder}")
        return

    print(f"\nFound {len(docx_files)} Word files in {input_folder}")
    for i, f in enumerate(docx_files, 1):
        print(f"  {i}. {f.name}")

    # Create output folder
    output_dir = input_folder / OUTPUT_DIR_NAME
    output_dir.mkdir(exist_ok=True)

    # Step 1: Merge to .docx
    merged_docx_path = output_dir / f"{OUTPUT_FILENAME}.docx"
    print("\nCombining documents...")
    merge_documents(
        docx_files, merged_docx_path,
        add_page_breaks=page_breaks,
        insert_toc=add_toc
    )
    print(f"  Saved merged document: {merged_docx_path}")

    # Step 2: PDF
    if "pdf" in formats:
        pdf_path = output_dir / f"{OUTPUT_FILENAME}.pdf"
        print("\nCreating PDF...")
        if convert_to_pdf(merged_docx_path, pdf_path):
            print(f"  Saved PDF: {pdf_path}")
        else:
            print("  PDF generation failed.")

    # Step 3: Markdown
    if "md" in formats:
        md_path = output_dir / f"{OUTPUT_FILENAME}.md"
        print("\nCreating Markdown...")
        if convert_to_markdown(merged_docx_path, md_path):
            print(f"  Saved Markdown: {md_path}")
        else:
            print("  Markdown generation failed.")

    print("\n✓ Complete!")

# ----------------------------------------------------------------------
# 5. Interactive CLI
# ----------------------------------------------------------------------
def interactive_mode():
    """Show a simple menu when no command-line arguments are given."""
    print("=" * 40)
    print(" Word Document Merger")
    print("=" * 40)

    # Try to use a folder picker dialog (tkinter) if available
    try:
        import tkinter as tk
        from tkinter import filedialog
        root = tk.Tk()
        root.withdraw()
        root.attributes('-topmost', True)
        folder = filedialog.askdirectory(title="Select folder containing .docx files")
        root.destroy()
        if not folder:
            print("No folder selected. Exiting.")
            return
        folder = Path(folder)
    except ImportError:
        # Fallback to manual input
        path_str = input("Folder path: ").strip().strip('"')
        folder = Path(path_str)
        if not folder.is_dir():
            print("Invalid folder.")
            return

    # Sorting
    sort = input("Sort by [n]ame or [d]ate? (default: name): ").strip().lower()
    sort_order = "date" if sort.startswith("d") else "name"

    # Formats
    print("\nWhat would you like to generate?")
    print("1. PDF only")
    print("2. Markdown only")
    print("3. PDF + Markdown")
    print("4. DOCX only (merged .docx)")
    print("5. All (DOCX + PDF + Markdown)")
    choice = input("Choice: ").strip()
    fmt_map = {
        "1": ["pdf"],
        "2": ["md"],
        "3": ["pdf", "md"],
        "4": ["docx"],
        "5": ["docx", "pdf", "md"]
    }
    formats = fmt_map.get(choice, ["docx", "pdf", "md"])

    # TOC and page breaks
    toc = input("\nAdd automatic Table of Contents? (y/n, default: y): ").strip().lower() != 'n'
    breaks = input("Add page break before each document? (y/n, default: y): ").strip().lower() != 'n'

    run_workflow(folder, formats, sort_order, add_toc=toc, page_breaks=breaks)

# ----------------------------------------------------------------------
# 6. Command-line argument parsing
# ----------------------------------------------------------------------
def main():
    parser = argparse.ArgumentParser(
        description="Merge .docx files and export as .docx, PDF, and/or Markdown."
    )
    parser.add_argument("folder", nargs="?", help="Folder containing .docx files")
    parser.add_argument("-o", "--output", help="Output directory (default: <input>/Output)")
    parser.add_argument("-f", "--formats", nargs="+", choices=["docx", "pdf", "md"],
                        default=["docx", "pdf", "md"],
                        help="Output formats to generate (default: all)")
    parser.add_argument("--sort", choices=["name", "date"], default="name",
                        help="Sorting order (default: name)")
    parser.add_argument("--no-toc", action="store_true",
                        help="Do not insert a Table of Contents")
    parser.add_argument("--no-page-breaks", action="store_true",
                        help="Do not insert page breaks between documents")
    parser.add_argument("--interactive", action="store_true",
                        help="Force interactive mode even with folder argument")

    args = parser.parse_args()

    # If no folder provided, enter interactive mode
    if not args.folder or args.interactive:
        interactive_mode()
        return

    folder = Path(args.folder)
    run_workflow(
        folder,
        formats=args.formats,
        sort_order=args.sort,
        add_toc=not args.no_toc,
        page_breaks=not args.no_page_breaks
    )

if __name__ == "__main__":
    main()